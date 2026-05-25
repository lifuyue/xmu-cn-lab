from __future__ import annotations

from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
REPO = ROOT.parents[1]
TEMPLATE = Path("/Users/lifuyue/Downloads/实验报告格式（20260328）.docx")
OUTPUT = REPO / "ex6" / "22920242203267_李富悦_实验8_应用层协议服务配置.docx"
REAL_SCREENSHOTS = ROOT / "artifacts" / "real_screenshots"
DOC_IMAGES = ROOT / "artifacts" / "real_screenshots_for_doc"


def set_value_after_first_tab(paragraph: Paragraph, value: str) -> None:
    """Replace only the fillable value in a metadata line, keeping template runs."""
    found_tab = False
    value_written = False
    for run in paragraph.runs:
        text = run.text
        if not found_tab and "\t" in text:
            before, _, _ = text.partition("\t")
            run.text = before + "\t" + value + "\t"
            found_tab = True
            value_written = True
            continue
        if found_tab:
            run.text = ""
    if not value_written:
        paragraph.add_run(value)


def replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace paragraph content while retaining the first run's direct formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def prepare_doc_images() -> dict[str, Path]:
    DOC_IMAGES.mkdir(parents=True, exist_ok=True)
    images: dict[str, Path] = {}
    for source in sorted(REAL_SCREENSHOTS.glob("[0-9][0-9]_*.png")):
        target = DOC_IMAGES / (source.stem + ".jpg")
        with Image.open(source) as image:
            rgb = image.convert("RGB")
            max_width = 1500
            if rgb.width > max_width:
                ratio = max_width / rgb.width
                rgb = rgb.resize((max_width, int(rgb.height * ratio)), Image.Resampling.LANCZOS)
            rgb.save(target, quality=82, optimize=True)
        images[source.stem] = target
    return images


def add_result_items(doc: Document, anchor: Paragraph, images: dict[str, Path]) -> None:
    items = [
        (
            "01_start_services",
            "启动服务与生成证书：生成本地 CA、服务器证书和 SSH 密钥，启动 DNS、HTTP、HTTPS、FTP、SSH、SMTP、POP3、IMAP 服务。",
        ),
        (
            "02_dns",
            "DNS 服务器：使用 dig 和 nslookup 查询 lab8.local 实验域名，返回 127.0.0.1。",
        ),
        (
            "03_http",
            "Web 服务器：访问 HTTP 首页，确认服务器返回 200 OK 和页面内容。",
        ),
        (
            "04_virtual_hosts",
            "虚拟主机：同一 IP 和端口下，site-a.lab8.local 与 site-b.lab8.local 返回不同页面。",
        ),
        (
            "05_auth_rate",
            "访问控制与流量控制：未认证访问返回 401，账号认证后返回 200，并演示限速下载。",
        ),
        (
            "06_https_cert",
            "安全站点与证书：查看 CA 签发的服务器证书，并通过 CA 根证书验证 HTTPS 访问。",
        ),
        (
            "07_ftp",
            "FTP 服务器：登录后完成上传、列目录和下载，权限策略拒绝新建目录。",
        ),
        (
            "08_smb",
            "SMB 共享：启动本地 SMB 共享并列出 LAB8 共享内容。",
        ),
        (
            "09_ssh",
            "SSH 服务器：使用本地密钥连接 2222 端口并远程执行命令。",
        ),
        (
            "10_mail",
            "SMTP、POP3 与 IMAP：通过 SMTP 投递测试邮件，并用 POP3、IMAP 读取邮件。",
        ),
    ]

    current = anchor
    for index, (stem, caption) in enumerate(items, 1):
        current = insert_paragraph_after(current, f"{index}. {caption}", anchor.style)
        image_paragraph = insert_paragraph_after(current, "", anchor.style)
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = image_paragraph.add_run()
        run.add_picture(str(images[stem]), width=Inches(5.65))
        current = insert_paragraph_after(image_paragraph, f"图 {index}　{caption}", anchor.style)
        current.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build() -> None:
    images = prepare_doc_images()
    doc = Document(str(TEMPLATE))
    paragraphs = doc.paragraphs

    set_value_after_first_tab(paragraphs[4], "实验八　应用层协议服务配置")
    set_value_after_first_tab(paragraphs[5], "软件工程2024级1班")
    set_value_after_first_tab(paragraphs[6], "李富悦")
    set_value_after_first_tab(paragraphs[7], "22920242203267")
    set_value_after_first_tab(paragraphs[8], "2026年5月25日")
    replace_paragraph_text(paragraphs[9], "")

    replace_paragraph_text(
        paragraphs[20],
        "掌握 DNS、HTTP/HTTPS、FTP、SSH、SMTP、POP3、IMAP、SMB 等应用层服务的基本配置方法；"
        "理解虚拟主机、访问控制、证书签发与 HTTPS 加密访问的配置要点；"
        "能够使用客户端命令验证服务监听、域名解析、身份认证、文件传输和邮件收发结果。",
    )
    replace_paragraph_text(
        paragraphs[22],
        "操作系统：macOS；实验目录：ex6/chapter8_app_services；工具：Python 3.9、OpenSSL、OpenSSH、curl、dig、nslookup、Impacket；"
        "服务端口：DNS 15353、HTTP 8080、HTTPS 8443、FTP 2121、SSH 2222、SMB 1445、SMTP 2525、POP3 8110、IMAP 8143。",
    )
    replace_paragraph_text(
        paragraphs[24],
        "以下截图均为本机真实 Terminal 窗口截图，截图时逐项启动服务并运行客户端验证命令。",
    )
    add_result_items(doc, paragraphs[24], images)

    replace_paragraph_text(
        paragraphs[25],
        "实验代码",
    )
    replace_paragraph_text(
        paragraphs[26],
        "本次实验的代码已上传于以下代码仓库：https://github.com/lifuyue/xmu-cn-lab/tree/main/ex6/chapter8_app_services。",
    )
    replace_paragraph_text(paragraphs[28], "第8章未给出独立课后思考题，本项填写为无。")
    replace_paragraph_text(
        paragraphs[30],
        "本次实验把手册中的应用层服务配置要求落实为可复现的本机服务验证。通过真实客户端命令验证 DNS 解析、"
        "Web 虚拟主机、HTTPS 证书、FTP 权限控制、SMB 共享、SSH 登录以及邮件收发，可以更清楚地区分服务启动、"
        "协议连通和应用功能验证之间的关系。",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    build()
